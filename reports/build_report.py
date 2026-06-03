# -*- coding: utf-8 -*-
"""Build the Google I/O 2026 briefing as a styled .docx for JOAT Kenya."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# JOAT brand-ish palette
NAVY = RGBColor(0x10, 0x18, 0x28)
GOLD = RGBColor(0xC8, 0x9B, 0x2C)
RED = RGBColor(0xC0, 0x2A, 0x2A)
GREY = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(0x20, 0x20, 0x20)

for lvl, sz, col in [("Heading 1", 17, NAVY), ("Heading 2", 13.5, RED), ("Heading 3", 11.5, NAVY)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell_text(cell, text, bold=False, color=None, size=9.5, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if white:
        run.font.color.rgb = WHITE
    elif color:
        run.font.color.rgb = color


def add_para(text, size=10.5, bold=False, italic=False, color=None, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_rich(parts, size=10.5, space_after=6):
    """parts = list of (text, bold) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold in parts:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
    return p


def add_bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.bold = True
        r.font.size = Pt(10.5)
        rest = p.add_run(text)
        rest.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
    return p


def callout(label, text, fill="FFF6E0", bar=GOLD, label_color=RED):
    """Single-cell shaded box used for 'Say safely' and caveats."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, fill)
    cell.width = Inches(6.6)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    rl = p.add_run(label + "  ")
    rl.font.bold = True
    rl.font.size = Pt(9.5)
    rl.font.color.rgb = label_color
    rt = p.add_run(text)
    rt.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def make_table(headers, rows, col_widths=None, header_fill="101828"):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], header_fill)
        set_cell_text(hdr[i], h, bold=True, white=True, size=9)
    for r in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(r):
            set_cell_text(cells[i], val, size=9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    return tbl


def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C89B2C")
    pbdr.append(bottom)
    pPr.append(pbdr)


# =====================================================================
# COVER
# =====================================================================
t = add_para("GOOGLE I/O 2026", size=26, bold=True, color=NAVY, space_after=2,
             align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("The 10 Headline Releases", size=18, bold=True, color=GOLD, space_after=2)
add_para("A tech-team + marketing briefing, mapped to the JOAT Kenya & Amare's Big Planet workflow",
         size=12, italic=True, color=GREY, space_after=10)
add_rich([("Prepared by: ", True), ("Tech Team    ", False),
          ("For: ", True), ("Tech peers + Marketing    ", False),
          ("Due: ", True), ("Monday", False)], size=10)
add_rich([("Date of analysis: ", True), ("1 June 2026    ", False),
          ("Event: ", True), ("Google I/O 2026 keynote (19 May 2026)", False)], size=10)
hr()

# =====================================================================
# SECTION 0
# =====================================================================
doc.add_heading("0. How to read this report (please read first)", level=1)
add_para("This is built for two audiences at once:")
add_bullet("what each release actually is, how it changes our build/ship workflow vs the old way, and where it fits next to the stack we already run in production (Anthropic Claude + OpenRouter, Cloudflare Workers, Supabase, React 19, Flutter/React Native, Blender, Roblox).",
           bold_lead="For your peers (tech): ")
add_bullet("plain-English value, price points, release dates, and a 'say-this-safely-on-stage' line per feature.",
           bold_lead="For marketing & the presentation: ")
add_rich([("Honesty rule we followed. ", True),
          ("Every claim was researched against official Google sources (blog.google, DeepMind, developer docs) and then independently fact-checked by a second pass. Some names in the original brief are imprecise or wrong - Google never shipped products called 'Asset Payment Protocol,' 'Smart Cart,' or a voice model called 'Neural Expressive.' We correct these in Section 2 so nobody quotes a fake product name on stage. Where a price or date couldn't be confirmed, we say so rather than guess.", False)])
callout("GOLDEN CAVEAT:",
        "Google said 'global' for many of these, but 'global' usually means subscriber/app access, not the developer API, and rarely names Africa. Only TWO of the ten are confirmed usable from Kenya today (Google Flow, and Google Finance Deep Search). Treat every other Kenya availability as UNCONFIRMED until tested.",
        fill="FDEAEA", bar=RED, label_color=RED)

# =====================================================================
# SECTION 1 - EXEC SUMMARY
# =====================================================================
doc.add_heading("1. Executive summary - the five things that matter", level=1)
exec_pts = [
    ("The theme was agents, not chatbots. ", "Nearly every release moves AI from 'answer my question' to 'go do the multi-step work and report back.' This validates JOAT's strategy and gives us new tools to plug in."),
    ("The 'cheap AI' era is ending. ", "Gemini 3.5 Flash is brilliant but 3-5x more expensive than the old Flash tiers. For an M-Pesa-context market, model choice is now a per-task cost decision - which is exactly why our OpenRouter 'route to cheapest-good-enough' approach is an advantage."),
    ("Most consumer features are US-first / subscription-gated. ", "That's a market gap. Google shows the world 'delegate a task to an agent' while leaving East Africa under-served. We can ship that value at local prices on the stack we already pay for."),
    ("The new money-makers for us are integration patterns, not the products. ", "MCP, UCP/AP2, object-level image editing, expressive voice are callable primitives we wrap around our Claude core - and can sell as new service lines."),
    ("Don't rip and replace - extend. ", "Claude + OpenRouter stays our brain layer. The Google releases are specialist modules: Gemini for voice and video, Nano Banana for image edits, grounding API for freshness. Owning the orchestration layer is the moat."),
]
for i, (lead, body) in enumerate(exec_pts, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    rn = p.add_run(f"{i}. ")
    rn.font.bold = True
    rl = p.add_run(lead)
    rl.font.bold = True
    rl.font.color.rgb = NAVY
    p.add_run(body)

# =====================================================================
# SECTION 2 - NAME CORRECTIONS
# =====================================================================
doc.add_heading("2. Name corrections before you build slides", level=1)
make_table(
    ["In the brief", "Correct name(s)", "What it really is"],
    [
        ["Universal Commerce Protocol (Asset Payment Protocol)", "Agent Payments Protocol (AP2)",
         "'Asset Payment Protocol' appears in ZERO Google sources. The payments layer is AP2, donated to the FIDO Alliance."],
        ["Universal Cart / Smart Cart", "Universal Cart",
         "'Smart Cart' is journalist shorthand, not a Google brand. The product is Universal Cart."],
        ["Neural expressive (assumed a voice model)", "'Neural Expressive' = the Gemini APP UI redesign. Voice engine = Gemini 3.1 Flash TTS.",
         "Do NOT present 'Neural Expressive' as a voice product - it's a design language. The expressive-voice tech is the separate Gemini 3.1 Flash TTS (15 Apr 2026)."],
        ["Gemini Search real-time stock data", "Three separate things: information agents, Google Finance Deep Search, Grounding with Google Search API",
         "No single product by that name, and NONE guarantee exchange-grade tick quotes - it's freshest web + market data, not a Bloomberg feed."],
        ["Antigravity 2.0", "Google Antigravity (agentic dev platform)",
         "Real and major. The '2.0' label is used in press but not stamped consistently in Google's own docs - say 'the new Antigravity platform' to be safe."],
    ],
    col_widths=[2.1, 2.0, 2.5],
)

# =====================================================================
# SECTION 3 - CHEAT SHEET
# =====================================================================
doc.add_page_break()
doc.add_heading("3. Presenter's cheat-sheet (one glance)", level=1)
make_table(
    ["#", "Feature", "One-line value", "Released", "Price (confirmed)", "Kenya now?"],
    [
        ["1", "Gemini Omni (Omni Flash)", "Any input -> physics-aware video you edit by chatting", "19 May 2026", "Bundled in AI plans; free in YouTube Shorts/Create; API price UNPUBLISHED", "App likely; API: no"],
        ["2", "Gemini 3.5 Flash", "Frontier coding/agentic at Flash speed", "19 May 2026", "$1.50 / $9.00 per 1M tokens; $0.15 cached", "API yes; residency unconfirmed"],
        ["3", "Antigravity platform", "Orchestrate fleets of coding agents", "19 May 2026", "Free preview; Ultra $100/$200; Gemini CLI sunsets 18 Jun", "Likely; verify"],
        ["4", "Gemini Spark", "24/7 personal agent doing work in background", "Beta ~25 May", "Requires AI Ultra ($100 or $200)", "NO - US-only beta"],
        ["5", "Real-time Search data", "Live cited data + 'watcher' agents", "I/O 2026; Finance now global", "API: 5,000 free/mo then $14/1k", "YES (Finance + API)"],
        ["6", "Universal Cart + UCP + AP2", "Open standard for agents to shop & pay", "UCP Jan 2026; Cart summer 2026 (US)", "No fees disclosed; protocols open-source", "NO - US-first, no M-Pesa"],
        ["7", "Gemini 3.1 Flash TTS", "Direct AI voices with emotion tags, 70+ langs", "15 Apr 2026", "$1 / $20 per 1M (~$0.03/min)", "YES via OpenRouter"],
        ["8", "Google Pics", "Object-aware image editor (edit one element)", "Testers now; Pro/Ultra summer 2026", "AI Pro $19.99; no Pics API (use Nano Banana API)", "App later; Nano Banana API yes"],
        ["9", "Google Flow Agent", "Agentic AI filmmaking (Veo 3.1 + Omni)", "19 May 2026", "Pro $19.99 (Veo Lite trial)/Ultra; Flow credits", "YES - Kenya supported (18+)"],
        ["10", "Gemini for Science", "AI co-scientist: lit review, hypotheses, discovery", "19 May 2026", "No price (register-interest/Labs)", "NO - gated"],
    ],
    col_widths=[0.25, 1.25, 1.95, 1.0, 1.85, 0.95],
)
callout("SUBSCRIPTION MAP (post-I/O, confirmed):",
        "Google AI Plus ~$7.99/mo  |  Pro ~$19.99/mo  |  Ultra: a NEW $100/mo tier AND the old top tier cut from $250 -> $200/mo (two Ultra tiers now). KES/M-Pesa pricing was NOT disclosed for any tier.",
        fill="EAF1FB", bar=NAVY, label_color=NAVY)

# =====================================================================
# SECTION 4 - DEEP DIVES
# =====================================================================
doc.add_page_break()
doc.add_heading("4. The ten features in depth", level=1)
add_para("Each entry: What it is -> vs the old way -> Tech-team uses -> Marketing uses -> JOAT integration -> Amare's integration -> Better alternatives -> Price -> Release/availability -> Agentic-era view -> Say-this-safely.",
         italic=True, color=GREY, size=9.5)

FEATURES = [
    {
        "title": "1) Gemini Omni (first model: Omni Flash)",
        "blocks": [
            ("What it is.", "A new family of generative 'world models.' You feed it references (image, text, video, or a voice sample) and it outputs video that respects real-world physics and keeps a character's face and voice consistent across scenes. Signature trick: conversational, in-place editing - say 'make it a thunderstorm' and it edits without re-rolling the whole clip. Every output carries a SynthID watermark."),
            ("vs the old way (Veo-style text-to-video).", "Before, fixing one detail meant regenerating the whole clip and praying the character matched - a slot machine. Omni makes iteration incremental and identity-stable. Physics grounding kills uncanny motion."),
            ("Tech-team uses.", "Prototype API-driven video once the API ships; build 'edit-by-instruction' features; rely on SynthID for provenance; pair with Gemini 3.5 Flash as the reasoning layer."),
            ("Marketing uses.", "Fast IG/TikTok/Shorts ad variants (free via YouTube Shorts Remix/Create); localize one master ad into many versions keeping the same brand mascot; concept boards without a full Blender render."),
            ("JOAT integration.", "'Rough cut in Omni, polish in Blender.' Generate alphabet/phonics micro-clips from existing Blender character refs for Drawalette. Platform-specific ad variants. Talent-spotlight promo clips for Majobo at Shorts cost."),
            ("Amare's integration.", "Feed existing Blender/Roblox character art + a voice sample so Amare stays visually and vocally consistent across every letter/song - more episodes per cycle; Blender reserved for flagship content. SynthID supports a kid-safe story for parents."),
            ("Better alternatives.", "Claude is NOT a competitor here (no video) - keep Claude/OpenRouter for reasoning and treat Omni as a media add-on. Sora (single OpenAI surface), Runway/Pika/Luma (cheaper for pure clips), ElevenLabs (standalone voice)."),
            ("Price.", "No standalone price. Bundled in AI Plus/Pro/Ultra; free in YouTube Shorts/Create. The Gemini API/Vertex price row was still BLANK as of late May - any per-token/per-second figures are third-party estimates, not Google numbers. Do not quote them."),
            ("Release/availability.", "Omni Flash rolling out from 19 May 2026 to subscribers + YouTube; developer/enterprise API 'in coming weeks' (not GA as of 1 Jun). 'Globally' for subscriber/YouTube; API Africa availability unconfirmed."),
            ("Agentic-era view.", "Omni collapses 'generate media' into one step an agent can drive with words + reference assets. Strategic stack: Claude/OpenRouter for orchestration, a Gemini media model for assets, SynthID for provenance. Wait for published API pricing before high-volume automation."),
        ],
        "say": "Gemini Omni Flash launched 19 May 2026, free in YouTube Shorts/Create and in our AI plans; developer API pricing isn't published yet - any token figures you've seen are estimates. DON'T call it 'real-time live A/V' (the 'real-time' part is multi-turn editing) and DON'T claim full any-to-any output (it's video-out only at launch).",
    },
    {
        "title": "2) Gemini 3.5 Flash",
        "blocks": [
            ("What it is.", "First model in the Gemini 3.5 family. Positioned as fast/cheap 'Flash,' but it BEATS the previous premium Gemini 3.1 Pro on several coding/agentic/multimodal benchmarks while running ~4x faster on output. 1M-token input context, multimodal input, text-only output, 'dynamic thinking' on by default, native agentic tooling (function calling, code execution, search-as-a-tool, MCP)."),
            ("vs the old way.", "You used to route simple work to cheap Flash and hard work to expensive Pro. Now one fast model covers both. The catch: it is NOT cheaper than old Flash - input $0.30 -> $1.50 (5x), output $2.50 -> $9.00 (~3.6x). Google's 'less than half the cost' claim is vs RIVAL frontier models, not Google's own old Flash."),
            ("Tech-team uses.", "Single agentic-coding model in IDE pipelines; whole-repo/large-PDF jobs on the 1M window; high-throughput batch where 4x speed cuts wall-clock; tool-heavy agents (MCP); cached-input pricing ($0.15/1M) for repeated prompts."),
            ("Marketing uses.", "Fast multimodal drafting (analyse a video brief -> copy); bulk social variants with structured output; long-context campaign analysis (brand guidelines as cached input); audio/video -> briefs."),
            ("JOAT integration.", "BioBiz notes + live translation (A/B vs our Claude/OpenRouter cost given the price jump); Majobo classification with structured JSON + cheap cached taxonomy prompt; Roblox-Lua/Blender-Python automation in agentic IDE; marketing content engine repurposing 3D films into cutdowns."),
            ("Amare's integration.", "Ingest episodes -> auto-generate phonics worksheets, summaries, bilingual prompts; low-latency in-Roblox NPC/tutor - BUT add strict kid-safety guardrails and verify data-handling before any live kid-facing use."),
            ("Better alternatives.", "Claude (ours) - safer known default for kids' content and tuned agentic flows. OpenRouter (ours) - route cheapest-good-enough per task; better for M-Pesa-priced workloads. Lighter/older Flash tiers for trivial jobs. ElevenLabs for any voiceover (3.5 Flash is text-only)."),
            ("Price.", "$1.50 / $9.00 per 1M input/output; $0.15/1M cached (plus ~$1.00/1M-tokens/hr cache storage). Non-global regions ~$1.65/$9.90."),
            ("Release/availability.", "GA 19 May 2026 across Gemini app, AI Mode, API, AI Studio, Vertex, Antigravity, Android Studio. Gemini 3.5 Pro expected ~a month later. Kenya data-residency unconfirmed."),
            ("Agentic-era view.", "The 'cheap' tier is now frontier-capable for long-horizon agents - but the 3-5x price rise over old Flash is the real headline. Make model choice a per-task cost decision (via OpenRouter), reserve 3.5 Flash for genuinely agentic/coding-heavy work, lean on cached-input pricing."),
        ],
        "say": "Gemini 3.5 Flash (GA 19 May) is $1.50/$9.00 per 1M tokens - about 40% cheaper than Gemini 3.1 Pro but 3-5x pricier than older Flash; benchmarks are Google's own, vs 3.1 Pro. Make NO Kenya data-residency claim.",
    },
    {
        "title": "3) Google Antigravity (agentic development platform)",
        "blocks": [
            ("What it is.", "Google's agent-first IDE, now a full platform: a desktop app orchestrating multiple coding agents in parallel, a Go-based CLI, an SDK to build/self-host your own agents, a Managed Agents API (agents reason, call tools, run code in an isolated Linux sandbox with persistent state), and an enterprise path. Powered by Gemini 3.5 Flash. Adds Scheduled (cron) Tasks for background automation."),
            ("vs the old way.", "v1 assisted one dev on one task. The new platform fans work out - sub-agents in parallel, cron agents overnight, one-click export from AI Studio to local, native Android/Firebase/Workspace hooks that collapse prototype-to-production glue."),
            ("Tech-team uses.", "Parallel sub-agents on a React 19 + Vite + TanStack feature (one scaffolds UI, one writes hooks, one writes tests); SDK-wrap a Cloudflare Workers/Supabase deploy agent on our infra; headless CLI agents in CI/CD; cron agents for nightly dependency audits, SEO/sitemap regen, schema-drift checks, error-log triage; isolated sandbox for untrusted codegen."),
            ("Marketing uses.", "Cron agents auto-draft weekly social and stage for review; Gemini-Audio live transcription to caption/repurpose client footage; fast landing-page/A/B variants on 3.5 Flash; SDK 'campaign agent' that turns a brief into a content calendar."),
            ("JOAT integration.", "Rebuild/extend BioBiz notes+translation on Managed Agents (A/B vs Claude first - translation quality is the product's core value); nightly Majobo classification/moderation via cron agent (keep Claude fallback); script the Blender->Roblox pipeline with parallel sub-agents + CLI; staged social campaigns with a human approval gate."),
            ("Amare's integration.", "Sub-agents + CLI to auto-generate Roblox Lua level code and batch-export character assets; cron tasks to draft YouTube metadata - human-in-the-loop for cultural/educational accuracy."),
            ("Better alternatives.", "Claude (ours, Claude Code/API) - benchmark Antigravity against it rather than replace. OpenAI Codex/AgentKit - provider-neutral. Cursor/Windsurf/Copilot Workspace - more mature day-to-day IDE help. OpenRouter (ours) keeps model choice open; Antigravity ties you to Gemini billing."),
            ("Price.", "Free public-preview 'Individual' tier ($0); AI Pro ~$20/mo; AI Ultra $100 (5x limits) / $200 (top tier, cut from $250). Per-token Antigravity pricing not disclosed. The '$100 bonus credits' promo (through 25 May) has expired."),
            ("Release/availability.", "Launched 19 May 2026, broadly available (individual tier in public preview). Legacy Gemini CLI sunsets 18 June 2026 - migrate to Antigravity CLI (paid Code Assist Standard/Enterprise + Cloud keep access). Kenya access unconfirmed; the 'EU GDPR endpoint' claim is unverified - don't state it."),
            ("Agentic-era view.", "This is leverage: parallel sub-agents + cron agents let a 5-person team operate like 10-15. The catch is vendor lock-in: our differentiator is a model-agnostic stack. Adopt selectively for orchestration/scheduling where Google's ecosystem already touches our work; keep Claude/OpenRouter as the reasoning core."),
        ],
        "say": "At I/O (19 May) Google launched the Antigravity agentic platform - desktop app, CLI, SDK, Managed Agents API on Gemini 3.5 Flash; free preview tier, Ultra $100/$200, legacy Gemini CLI sunsets 18 June. Avoid the EU/GDPR endpoint and Kenya claims; frame AI Pro $20 as 'confirm on antigravity.google.'",
    },
    {
        "title": "4) Gemini Spark - 24/7 personal AI agent",
        "blocks": [
            ("What it is.", "An always-on agent on Google Cloud VMs that runs long/recurring tasks while your devices are off. Built on Gemini 3.5 + the Antigravity harness. Pulls context from Workspace (Gmail, Docs, Slides, Sheets) and the web, drafts work, and asks before high-stakes actions. Connects to third-party tools via MCP (Canva, OpenTable, Instacart at launch)."),
            ("vs the old way.", "The old Gemini app was reactive and tied up your device. Spark inverts it to delegation: hand off a multi-step job, it runs to completion in the background and reports back - including recurring/scheduled jobs."),
            ("Tech-team uses.", "Inbox/ops triage -> daily digest; recurring back-office automation (statements, invoices, M-Pesa receipts -> anomaly flags); async document generation; validate MCP as the integration standard - a signal to expose BioBiz/Majobo as MCP servers. NOTE: it's a closed consumer product, no API/SDK - you can't embed it."),
            ("Marketing uses.", "Standing social-content agent (email a brief -> drafts); always-on lead monitoring -> digest; recurring performance reports; Canva MCP for first-draft creative. Reality check: US-only + Ultra-gated = the Nairobi studio can't use it day-to-day yet."),
            ("JOAT integration (pattern-replication, since Spark is US-only).", "Expose BioBiz notes+translation as an MCP server so any agent (Spark AND Claude) can call it - turning BioBiz into agent-accessible infrastructure. Build a Spark-style recurring agent on our Claude/OpenRouter stack to auto-classify Majobo listings and email shortlists, at M-Pesa economics."),
            ("Amare's integration.", "A background agent (on our stack) monitors YouTube comments/submissions -> weekly digest of episode ideas + draft localized briefs celebrating African themes."),
            ("Better alternatives.", "Claude (ours) - strictly better for anything we need to embed (Spark has no SDK, US-only). OpenAI agents/Operator - programmatic + more regions. Self-hosted MCP + Claude orchestration honestly beats Spark for our purposes today, in-region, on infra we already pay for."),
            ("Price.", "Requires AI Ultra ($100 or $200 tier). No per-task price. CORRECTION: this was NOT a '$250 -> $100 price cut.' Google ADDED a new $100 Ultra tier AND separately cut the old top tier $250 -> $200. Spark works on both Ultra tiers."),
            ("Release/availability.", "Announced 19 May; beta to US AI Ultra subscribers ~25 May. Email/chat control and Chrome agentic browsing come over summer 2026 (at launch you use it via the Gemini app). No Kenya/international timeline."),
            ("Agentic-era view.", "Spark is the clearest 'chat -> delegated persistent agent over email/MCP' signal. Takeaways: (1) MCP is the de-facto integration fabric - expose BioBiz/Majobo as MCP servers; (2) the winning UX is 'delegate -> get result later,' ship that in our products; (3) the US-only/$100 gate leaves a real opening to deliver the same value at M-Pesa-friendly prices. Watch it as a design reference, not a dependency."),
        ],
        "say": "Gemini Spark (announced 19 May) is a cloud 24/7 personal agent, in US-only beta for AI Ultra ($100 or $200); email/chat and Chrome browsing come over summer 2026; no Kenya availability announced. DON'T claim a flat '$250->$100 cut' or that you can email it today.",
    },
    {
        "title": "5) Real-time data in Search (information agents | Finance Deep Search | Grounding API)",
        "blocks": [
            ("What it is.", "Not one product - THREE real things: (a) information agents in Search that continuously watch Google's freshest finance/sports/shopping/news data and ping you on changes; (b) Google Finance Deep Search - natural-language market questions with cited answers and live Kalshi/Polymarket odds; (c) Grounding with Google Search in the Gemini API - attach real-time, cited web data to your app. Gemini 3.5 Flash is now the default AI-Mode model globally, free."),
            ("vs the old way.", "Answers were bounded by training cutoff (stale) or needed manual searching. Now: real-time grounding closes the freshness gap, and information agents flip pull -> push (it watches and alerts). For devs, search-grounding is a billable API tool instead of bolt-on scraping."),
            ("Tech-team uses.", "Add grounding to a Gemini call for current, cited answers; build a market/news monitor in a Cloudflare Worker on a cron -> Supabase alerts; show source links in-app for trust; prototype an 'information-agent'-style watcher via API now."),
            ("Marketing uses.", "Same-day cited trend research; timely fact-grounded hooks; competitor/hashtag monitoring; current local context (events/weather/prices) for geo-targeted Kenya copy; claim-checking before publishing."),
            ("JOAT integration.", "BioBiz 'company brief' - after a meeting, pull fresh public info about a scanned contact's company and append to the AI notes. Majobo - enrich listings with market-rate/demand signals; keep classification aware of trending skills. Marketing pipeline anchored to real-time, citable sources. API watcher monitoring Roblox platform/asset-store changes."),
            ("Amare's integration.", "A grounded 'cultural freshness' helper pulling current, source-cited facts about African festivals/animals/places so each episode and Roblox level is accurate - citations the education team can fact-check before publishing."),
            ("Better alternatives.", "Claude + web search (ours) - the better default for our existing flows; Gemini wins only where you want Google index freshness or native Finance/prediction-market data. Perplexity API - cited Q&A, often cheaper. Polygon/Alpha Vantage/Finnhub - for exchange-grade quotes (LLM grounding is web-fresh, not tick-accurate). Kalshi/Polymarket APIs direct for odds."),
            ("Price.", "API grounding: 5,000 grounded prompts/month free on Gemini 3.x, then $14 per 1,000 queries (Gemini 2.5: $35/1k). Consumer Finance/agent features gated to AI Pro/Ultra; no standalone price."),
            ("Release/availability.", "Information agents: summer 2026, AI Pro/Ultra first (US-leaning). Google Finance Deep Search is now GLOBALLY available (May 2026) and usable from Kenya. API grounding: global. One of the few pieces with a clear Kenya path today."),
            ("Agentic-era view.", "The shift is push over pull - apps that watch the live world and act. The near-term lever for us is the grounding API: a cheap, citation-backed freshness primitive alongside our Claude stack. The consumer Search features are mostly US/Pro-gated, so the value for us is in the API, not the app."),
        ],
        "say": "Present as THREE separate items. 'Google announced information agents (summer 2026, Pro/Ultra first) and the Grounding API for developers; Google Finance Deep Search is already globally available and usable from Kenya.' Don't imply tick-grade stock quotes or one bundled product.",
    },
    {
        "title": "6) Universal Cart + Universal Commerce Protocol (UCP) + Agent Payments Protocol (AP2)",
        "blocks": [
            ("What it is.", "The open plumbing for agents that shop and pay. UCP (open standard, Jan 2026): merchants expose a /.well-known/ucp endpoint; agents call standard ops over MCP/JSON-RPC. Universal Cart: one Gemini-powered cart that follows you across Search, Gemini, YouTube, Gmail. AP2: agents pay within preset limits using cryptographically signed mandates (W3C Verifiable Credentials) for a tamper-proof audit trail; donated to the FIDO Alliance."),
            ("vs the old way.", "Agent checkout used to mean scraping stored card credentials, and networks couldn't tell a human from an agent or verify scope. The double-signed CartMandate + PaymentMandate fixes that. For devs, one discovery endpoint replaces bespoke per-merchant integrations. Build a checkout/payments agent once against an open standard."),
            ("Tech-team uses.", "Prototype a UCP merchant agent against the open ucp-sdk + google-adk codelab; stand up a UCP discovery endpoint on Cloudflare Workers/Supabase for a client store; evaluate AP2 mandate-signing as a pattern for auditable agent actions generally; wire AP2/UCP into our Claude stack via MCP."),
            ("Marketing uses.", "Make retail clients' products appear in Universal Cart + qualify for Direct Offers/Shopping ads; benchmark 'share of voice' on AI shopping surfaces; rewrite feeds with UCP 'conversational attributes'; pitch 'be findable by AI agents' as a new SEO/AEO service line; BNPL-aware creative (Affirm/Klarna in Google Pay)."),
            ("JOAT integration.", "Majobo - model the 1,000+ listings behind a UCP-style endpoint + AP2-style signed mandates so an agent can book/pay for a vetted gig within a budget cap. BioBiz - after a meeting, agent assembles a cross-vendor cart and pays under a per-meeting cap. Drawalette - UCP feeds so our store titles surface in Universal Cart. Marketing studio: a 'get your store agent-ready' service for SME clients."),
            ("Amare's integration.", "Expose Amare's merch/content (storybooks, plush, Roblox passes/bundles) via a UCP catalog so a parent can ask an agent to 'buy the Amare phonics bundle' - with an AP2 spending cap as built-in parental control."),
            ("Better alternatives.", "OpenAI/Stripe 'Agentic Commerce Protocol' + ChatGPT Instant Checkout - better if clients want to sell inside ChatGPT. Claude (ours) + MCP - the agent brain; UCP/AP2 are a commerce layer, complementary. Shopify native tooling if a client's already on it. M-Pesa/Daraja + Flutterwave/Pesapal - the ONLY realistic Kenya rails today; AP2 is US-first, card-centric, no mobile-money shown. Visa Intelligent Commerce / Mastercard Agent Pay as alternates."),
            ("Price.", "No fees disclosed (present as 'no fees announced,' not 'free'). Protocols + SDKs are open-source."),
            ("Release/availability.", "UCP standard Jan 2026; Universal Cart US, Search+Gemini, summer 2026 (YouTube/Gmail later); UCP checkout expanding to Canada/Australia (UK later); AP2 to Google products 'starting with Spark.' No Kenya/Africa availability; India only for one Merchant Center tool."),
            ("Agentic-era view.", "Commerce just got a machine-callable, verifiable interface. Our edge: we already build agents (Claude + MCP). Treat 'agent-readiness' as a productized service and bake mandate-style auditable autonomy into BioBiz/Majobo now, so when the rails reach East Africa we've shipped the patterns. Hedge against US-first/card-first dependency by keeping M-Pesa rails first-class."),
        ],
        "say": "Use the real names - Universal Cart, UCP, AP2 - never 'Smart Cart' or 'Asset Payment Protocol.' 'US-first rollout (summer 2026) of open, fee-undisclosed agentic-commerce standards; no Kenya/Africa availability announced; no M-Pesa support shown.'",
    },
    {
        "title": "7) The real expressive-voice release: Gemini 3.1 Flash TTS (brief's 'Neural expressive')",
        "blocks": [
            ("What it is.", "CORRECTION FIRST: 'Neural Expressive' is the Gemini APP design-language redesign (fluid animation, gradients, haptics, magazine-style answers) - a UI, not a voice model. The ACTUAL expressive-speech engine is Gemini 3.1 Flash TTS (15 Apr 2026): text->high-fidelity speech you 'direct like a film actor' with 200+ bracketed audio tags ([whispering], [excited], [sigh]) across 70+ languages, ~28-30 named voices, native multi-speaker dialogue, reusable Audio Profiles. All output SynthID-watermarked."),
            ("vs the old way.", "Over Gemini 2.5 TTS: instead of just picking a voice, you embed stage directions inline, get multi-character dialogue in one call (huge for podcasts/character work), and consistent character voices across renders. From flat narration -> scripted, emotionally-directed, multi-voice audio without stitching calls."),
            ("Tech-team uses.", "Wire it into our Cloudflare Workers/Supabase backend via the Gemini API OR our existing OpenRouter account (minimal new infra); script multi-speaker dialogue in one call; cache Audio Profiles for recurring voices; adopt the Neural-Expressive UI pattern as a design reference for our React 19 assistant UIs; SynthID for compliance."),
            ("Marketing uses.", "Expressive multi-language ad voiceovers ~$0.03/min (no booking talent); two-voice podcast-style social clips; localize one script into 70+ languages with emotion; Google Vids built-in voiceovers; A/B read styles by swapping tags instead of re-recording."),
            ("JOAT integration.", "BioBiz - speak the translated output back naturally and read summaries aloud -> turns text translation into a voice-assisted networking tool. Drawalette - warm narrated letter sounds + encouragement ([excited] Great job!). Blender/Roblox - multi-character NPC/explainer dialogue, cutting voice-acting cost. Majobo - spoken listing playback / voice search for low-literacy, mobile-first users."),
            ("Amare's integration.", "Voice phonics + African-culture characters across YouTube and the Roblox game; multi-speaker dialogue + reusable Audio Profiles keep each character's distinct voice; emotion tags make letter-sound teaching warm for ages 1-8; rapid localization - QA Swahili quality (per-language tiers unverified)."),
            ("Better alternatives.", "ElevenLabs - still the benchmark for cinematic emotional TTS / voice cloning / dubbing; better for premium brand films and owned voices. OpenAI Realtime+TTS - strong for low-latency conversational voice. Claude (ours) does NOT do TTS - Gemini TTS complements, doesn't replace, our LLM layer. Open-source XTTS-style for high-volume Roblox NPC audio. Classic Google Cloud TTS for plain narration."),
            ("Price.", "$1.00 / $20.00 per 1M input/output tokens (audio ~25 tokens/sec => ~$0.03 per 60-second voiceover) - this is on Google's official pricing page, with a free preview tier. Available via OpenRouter too."),
            ("Release/availability.", "TTS launched 15 Apr 2026 (public preview) via Gemini API, AI Studio, Vertex, Google Vids, and OpenRouter. Neural-Expressive UI rolling out from 19 May 2026. Kenya reachable via API/OpenRouter; Swahili supported, quality unverified."),
            ("Agentic-era view.", "Voice is becoming the OUTPUT layer of agents, not a separate product. Our existing Claude/OpenRouter agents can now speak with emotion, in local languages, at ~$0.03/min. Treat it as a modular voice tool (Claude for brains, Gemini for voice) - no platform lock-in. The Neural-Expressive UI is a north-star for presenting our own agent outputs: structured, multimedia, less text-wall."),
        ],
        "say": "'Neural Expressive' is Google's new Gemini app design language (from 19 May) - a UI, NOT a voice feature. The expressive-voice engine is the separate Gemini 3.1 Flash TTS (preview since 15 Apr): 200+ audio tags, 70+ languages, $1/$20 per 1M tokens (~$0.03/min), reachable via OpenRouter. Flag Kenya/Swahili quality as not-yet-verified.",
    },
    {
        "title": "8) Google Pics",
        "blocks": [
            ("What it is.", "A standalone web app (NOT Google Photos, NOT a rebrand) that treats every element of an image as a separately editable object. Click an element and leave a text comment (like commenting in Docs) to recolor, swap, move/resize, or edit & translate on-image text while preserving font/layout. Works on your uploaded photos too. Runs on Google's Nano Banana image model, SynthID-watermarked, integrated with Slides/Drive. Aimed at non-designers."),
            ("vs the old way.", "Traditional generators (including what we wire via OpenRouter) are 'flat' - change one detail and you re-roll, often losing what you liked. Pics keeps an object-level representation so one element changes deterministically - a slot-machine workflow becomes a layered-editor workflow driven by plain language. On-image text translation that keeps typography is a step beyond generic generators."),
            ("Tech-team uses.", "The Pics APP has no API. But the underlying Nano Banana models ARE available via the Gemini API (AI Studio/Vertex) - that's the dev path. Prototype comment-to-edit UX by calling Nano Banana for iterative edits; use its on-image text + translation to auto-localize graphics (English<->Swahili banners); evaluate SynthID handling; use the Pics app as a fast internal slide/Drive asset tool."),
            ("Marketing uses.", "Spin campaign creatives and edit one element per platform (swap CTA, recolor product) without regenerating; localize one master ad into English+Swahili keeping exact layout/font; flyers/infographics for SME clients with no designer; live collaborative canvas with clients; clone-and-tweak A/B variants."),
            ("JOAT integration.", "Marketing studio's day-to-day fast-creative tool (Pro seat) - localize one master into many platform/language variants. BioBiz - extend its translation story to visuals via Nano Banana API. Drawalette/Blender - iterate 2D concept art and alphabet-letter illustrations before committing to Blender. Majobo - auto-generate + locally-translate listing/category banners at scale."),
            ("Amare's integration.", "Mass-produce phonics/alphabet graphics - generate a consistent African-culture character, then use object segmentation to drop it into 26 letter scenes, translating on-image labels English<->Swahili - feeding YouTube thumbnails and Roblox signage without redrawing each variant."),
            ("Better alternatives.", "Claude (ours, Claude Design) - explicitly named a Pics competitor; staying on Claude avoids a new vendor. OpenAI GPT-image via OpenRouter (ours) - incumbent for programmatic generation. Canva - friendlier/cheaper for non-technical staff. Adobe Firefly/Photoshop for deep pro control. Honest verdict: the Pics app doesn't beat our Claude+OpenRouter production stack (no API) - evaluate Nano Banana via the Gemini API head-to-head on cost/quality vs our current OpenRouter image models."),
            ("Price.", "No standalone price; in AI Pro ($19.99) and Ultra ($100/$200) - Plus ($7.99) does NOT get Pics. Workspace business customers get preview. KES/M-Pesa not disclosed."),
            ("Release/availability.", "Trusted testers now; summer 2026 to AI Pro/Ultra 'globally' + Workspace preview. Web-app first. No Pics API (use Nano Banana API). Kenya timing unconfirmed; Nano Banana API usable now."),
            ("Agentic-era view.", "Google is pushing image editing toward a structured, object-addressable canvas - exactly what an agent needs to make precise, verifiable visual edits programmatically. The app is human-in-the-loop, but the Nano Banana API gives our agents a 'change just this element' tool -> automated localization and creative-variant pipelines. Don't depend on the consumer app; wrap Nano Banana as an agent tool alongside Claude/OpenRouter."),
        ],
        "say": "Google Pics (announced 19 May) is a Workspace-native AI image editor on Google's latest Nano Banana model - testers now, summer 2026 to AI Pro ($19.99)/Ultra. It's NOT Google Photos, and there's NO Pics API - devs use the Nano Banana API. State Ultra as $100/$200; don't claim Kenya/M-Pesa availability.",
    },
    {
        "title": "9) Google Flow Agent (agentic AI filmmaking)",
        "blocks": [
            ("What it is.", "Flow is Google's AI 'creative studio' (built on Veo 3.1, merged with Whisk/ImageFX). The new Flow Agent takes multi-step creative tasks in one session: plan a sequence, batch-generate scene variations, suggest plot/dialogue, run batch edits, organize assets - conversationally. Alongside it: Gemini Omni Flash (character/voice consistency across scenes) and Flow Tools (vibe-code custom editor/resizer/shader presets, no code)."),
            ("vs the old way.", "Old Flow ran one prompt at a time - fire, eyeball, re-prompt, hand-stitch. The Agent holds project context, plans multi-step, fires batched variations, iterates conversationally -> from 'operate the tool prompt-by-prompt' to 'brief a creative assistant and curate.' Omni adds character/voice consistency (the classic AI-video failure where the face changes every shot)."),
            ("Tech-team uses.", "Rapid-prototype explainers/product demos (brief once -> curate batched variations); watch for the Omni/Veo API to wire agentic video into our pipelines; Flow Tools presets to standardize aspect-ratio/brand-look; reproducible storyboard-to-render step. RELIABILITY: reported high prompt-failure/silent-error rates mean R&D/spike work, not a hard-deadline path yet."),
            ("Marketing uses.", "Multiple ad-creative variations in one session (batch variation is the core strength); auto-resize one master into Reels/Shorts/feed via a saved Flow Tool; conversational iteration without re-briefing an editor; Flow Music for branded music clips/jingles; brand-character consistency across a campaign via Omni."),
            ("JOAT integration.", "Insert Flow as pre-viz/storyboard + rough-cut ahead of Blender, then finish hero assets in Blender. BioBiz - batch localized app-promo videos (English + Kenyan-context) with a consistent spokesperson. Majobo - explainer/onboarding videos + seasonal ad variations. Drawalette - promo trailers/teasers (NOT in-product kids' likenesses - see caveat). Flow Tools presets for consistent multi-format ad cuts."),
            ("Amare's integration.", "Storyboard/rough-cut episode concepts - generate several African-culture/phonics scene variations to test pacing, then hand the locked concept to the Blender pipeline for final, brand-safe production; Flow Music for phonics song hooks. IMPORTANT: Flow's moderation aggressively blocks minor-related generation, so it's a pre-viz/trailer/music-draft tool for Amare's - NOT a generator of the actual child-facing characters."),
            ("Better alternatives.", "Sora 2 (may win on some physics/realism). Runway Gen-4 (finer manual control/motion brush). ElevenLabs (superior standalone TTS/dubbing - pair rather than rely on Veo audio). Claude+OpenRouter (ours) - Flow does NOT replace our stack (transcription/translation/scripts/agentic automation); keep Claude for scripts, add Flow as an additive video module only if the Veo/Omni API materializes. Blender (in-house) - still better for final kids' assets, ownership, and avoiding moderation false-positives."),
            ("Price.", "No standalone price; gated by AI tiers. AI Pro $19.99/mo (limited Veo 3.1 Lite trial, 5TB, $10/mo Cloud credit); Ultra $100 (5x) / $200 (20x, cut from $250). Flow consumes 'Flow credits' - reportedly deducted even on failed generations (user reports, not official). Kenya/M-Pesa cost not disclosed."),
            ("Release/availability.", "19 May 2026, live to all Flow users globally - and Kenya is explicitly listed as a supported country (190+ countries; 18+ only). Android app in beta; iOS later."),
            ("Agentic-era view.", "I/O's throughline applied to creative production: the human moves from operator -> director/curator. Risk: generating 16 variations relocates work to curation rather than removing it, and reported failure rates show agentic creative isn't deadline-grade yet. Treat Flow as one agent in a portfolio (video) orchestrated by our reasoning layer, with Blender/human finishing as the quality backstop."),
        ],
        "say": "Google Flow Agent (19 May 2026) is live to all Flow users INCLUDING Kenya, gated by AI tiers (Pro $19.99, Ultra $100/$200). Flag 'credits charged on failed generations' as a reported user concern, not official, and note the strict minor-content moderation for kids' work.",
    },
    {
        "title": "10) Gemini for Science",
        "blocks": [
            ("What it is.", "A DeepMind/Google Research program of experimental 'AI co-scientist' tools, opening gradually via Google Labs: Literature Insights (NotebookLM engine), Hypothesis Generation (multi-agent AI Co-Scientist 'idea tournament'), Computational Discovery (AlphaEvolve + ERA - generate/score thousands of code variants in parallel), plus Science Skills wiring Gemini into 30+ life-science databases inside Antigravity. Aimed at PhD-level researchers. Backed by two peer-reviewed Nature papers (19 May 2026) - ERA reportedly beat the CDC's COVID hospitalization forecasting ensemble."),
            ("vs the old way.", "Co-Scientist and AlphaEvolve were research demos/papers, 'not publicly available.' This productizes them as Labs tools with Nature-grade validation and a unified agentic workbench. From 'read paper -> hand-wire scripts -> hand-tune for weeks' to an agent running analyses in minutes and thousands of parallel variants automatically."),
            ("Tech-team uses.", "Use Literature Insights/Deep Research as a faster competitive-tech scan before building; study the Co-Scientist 'idea tournament' (generate->debate->score->refine) as an architecture reference for our own agentic automation on Claude+OpenRouter; borrow parallel variation-and-score for internal eval harnesses; study Science Skills' skill/tool-registry pattern."),
            ("Marketing uses.", "Repurpose NotebookLM-style auto infographics/decks/audio-video overviews as a template for fast thought-leadership; position JOAT as AI-fluent with a credible 'hype vs fact' explainer (like this report); Deep Research for data-backed pitch briefs; myth-vs-fact carousels to attract research/edtech/health-tech clients; the Nature/ERA-beats-CDC angle as a B2B credibility hook."),
            ("JOAT integration.", "BioBiz - Co-Scientist's clickable-citation + multi-agent verification pattern makes AI meeting notes more trustworthy; not a replacement for our Claude transcription. Majobo - the scoring loop to improve listing classification/ranking. Drawalette - Deep Research/Literature Insights to pull evidence-based phonics/early-literacy research informing new levels. NotebookLM-style overviews to speed client deliverables."),
            ("Amare's integration.", "Ground content in peer-reviewed early-literacy + African-culture sources, turn findings into NotebookLM-style audio/video overviews feeding scriptwriting and Roblox level design - making the literacy/STEM-curiosity outcomes demonstrably evidence-based for parents and education partners."),
            ("Better alternatives.", "Claude (ours) - better day-to-day for general agentic automation/content/summarization; Gemini for Science is research-domain-specific and gated. OpenRouter (ours) keeps us model-agnostic/cheaper. OpenAI Deep Research - comparable, more openly available API. AlphaFold/specialized tools still beat a general agent on their own domain. ElevenLabs beats NotebookLM audio for Amare's narration."),
            ("Price.", "No pricing disclosed - register-interest/Labs access (don't call it 'free,' that's an inference); enterprise via Google Cloud private preview (partners include BASF, Bayer Crop Science, Daiichi Sankyo, Klarna)."),
            ("Release/availability.", "Announced 19 May 2026; Nature papers same day; access opening gradually from May 2026 via labs.google/science (trusted-tester) + Cloud private preview. Not GA. Kenya unconfirmed and likely limited at launch."),
            ("Agentic-era view.", "The clearest signal that the frontier is shifting to autonomous, multi-agent systems that plan, call tools, and verify their own work. The takeaway isn't to adopt the gated product; it's to internalize the architecture: idea-tournament hypothesis loops, skill/tool registries bound to trusted data, parallel variation-and-score, citation-backed outputs. Apply these on our Claude+OpenRouter foundation. NOTE: per MIT Tech Review, AI-science results are largely lab-validated, not clinic-validated - present benchmarks as reported, not settled."),
        ],
        "say": "Gemini for Science (19 May 2026) is three experimental Labs tools (Co-Scientist, AlphaEvolve+ERA, NotebookLM-based Literature Insights), backed by two Nature papers the same day. Register-interest access; enterprise in private preview; no consumer price; Kenya unconfirmed. Don't call it 'free' or generally available.",
    },
]

for idx, feat in enumerate(FEATURES):
    if idx > 0:
        doc.add_page_break()
    doc.add_heading(feat["title"], level=2)
    for lead, body in feat["blocks"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        rl = p.add_run(lead + " ")
        rl.font.bold = True
        rl.font.color.rgb = NAVY
        rl.font.size = Pt(10.5)
        rb = p.add_run(body)
        rb.font.size = Pt(10.5)
    callout("SAY SAFELY:", feat["say"], fill="FFF6E0", bar=GOLD, label_color=RED)

# =====================================================================
# SECTION 5 - WHAT IT MEANS FOR OUR WORKFLOW
# =====================================================================
doc.add_page_break()
doc.add_heading("5. What this means for OUR workflow", level=1)

doc.add_heading("5.1 The honest adoption scorecard", level=3)
make_table(
    ["Verdict", "Features", "Why"],
    [
        ["ADOPT / PILOT NOW", "Gemini 3.1 Flash TTS (via OpenRouter), Grounding API + Finance Deep Search, Google Flow (Kenya-supported), Nano Banana API (Pics' engine)",
         "Reachable from Kenya today, clear ROI, plug into existing stack/products."],
        ["TEST vs OUR STACK", "Gemini 3.5 Flash, Antigravity platform",
         "Powerful but pricier/lock-in risk - A/B vs Claude+OpenRouter on cost-per-task before committing."],
        ["WATCH (design reference)", "Gemini Spark, Gemini Omni (API), Gemini for Science",
         "Not usable in-region yet. Copy the PATTERNS (delegate-by-email, MCP, idea-tournaments) on our own stack."],
        ["PRODUCTIZE as a service", "UCP/AP2 'agent-ready commerce'",
         "Not a Kenya consumer play yet, but a new revenue line: build clients' UCP endpoints + auditable payments now."],
    ],
    col_widths=[1.4, 2.9, 2.4],
)

doc.add_heading("5.2 Six concrete moves (a 30/60/90 sketch)", level=3)
add_para("Next 30 days (no/low cost):", bold=True, color=RED, space_after=3)
add_bullet("Wire Gemini 3.1 Flash TTS into BioBiz (speak translations aloud) and Drawalette (narrated letters) via our existing OpenRouter account - ~$0.03/min, minimal infra.")
add_bullet("Add the Grounding-with-Search API (5,000 free/mo) to one feature - e.g. a BioBiz 'company brief' after a meeting.")
add_bullet("Spin up Google Flow (Kenya-supported) for marketing pre-viz/ad-variant tests - keep Blender as the finisher.")
add_para("Next 60 days (test + decide):", bold=True, color=RED, space_after=3)
add_bullet("Run a cost-per-task bake-off: Gemini 3.5 Flash vs our Claude/OpenRouter routing on Majobo classification and BioBiz translation.")
add_bullet("Expose BioBiz + Majobo as MCP servers - the single highest-leverage move; makes our products callable by any agent and is the foundation for everything agentic.")
add_para("Next 90 days (new revenue):", bold=True, color=RED, space_after=3)
add_bullet("Package 'Get your store agent-ready' (UCP endpoint on Cloudflare Workers/Supabase + AP2-style auditable payments + Demand-Gen campaigns) as a service line for SME clients - keeping M-Pesa rails first-class.")

doc.add_heading("5.3 The strategic thread for the 'AI agentic era' slide", level=3)
add_para("Google just spent its biggest keynote telling the world that the future is delegated, persistent, tool-using agents - and pricing the cheap tier up while gating the best consumer agents to the US. For JOAT that's a double opportunity:")
add_bullet("We already live where they're pointing. Claude + OpenRouter + MCP means we own the orchestration layer - the actual moat. Google's releases are modules we slot in (voice, video, images, grounding), not a stack we must adopt wholesale.")
add_bullet("The gap they left is our market. Spark-style 'delegate-and-get-results' UX, agent-ready commerce, and expressive local-language voice are all things we can ship at M-Pesa-friendly prices for East Africa while the big consumer versions stay US-locked. Being agent-native, model-agnostic, and local-first is the whole strategy - and I/O 2026 validated it.")

# =====================================================================
# SECTION 6 - RISK REGISTER
# =====================================================================
doc.add_heading("6. Risk register (present with credibility, not hype)", level=1)
risks = [
    ("Wrong names sink credibility. ", "Use AP2 / Universal Cart / Gemini 3.1 Flash TTS - never 'Asset Payment Protocol,' 'Smart Cart,' or 'Neural Expressive voice.'"),
    ("Don't quote unpublished prices as fact - ", "Omni API pricing is blank; any per-token figures are third-party estimates."),
    ("'Global' is not Kenya. ", "Only Flow and Finance Deep Search are confirmed in-region. Flag the rest as unconfirmed."),
    ("The 'cheap AI' story flipped - ", "Gemini 3.5 Flash is 3-5x pricier than old Flash. Lead with that honesty for a budget audience."),
    ("Kids' content needs care - ", "Flow blocks minor-related generation; any kid-facing live AI needs strict guardrails and data-handling checks. Keep Blender as the source of truth for actual Amare's/Drawalette characters."),
    ("Lock-in is the trap - ", "adopt Google modules selectively; keep Claude/OpenRouter as the brain."),
]
for lead, body in risks:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    rl = p.add_run(lead)
    rl.font.bold = True
    p.add_run(body)

# =====================================================================
# SECTION 7 - SOURCES
# =====================================================================
doc.add_heading("7. Sources & method", level=1)
add_para("Every claim was cross-checked against official Google sources (blog.google, DeepMind, Google Cloud, developer docs, official pricing pages) plus credible tech press (TechCrunch, The Verge / 9to5Google, The Next Web, MIT Technology Review), and then run through a second adversarial fact-check pass to catch hallucinated names, prices, and dates. Key primary sources include the blog.google I/O 2026 announcement hub, deepmind.google/models, ai.google.dev/gemini-api/docs/pricing, antigravity.google, ap2-protocol.org, and ai.google/gemini-for-science.")
add_para("Prepared for internal education + the standalone presentation. Happy to turn any section into slides or a one-pager. - Tech Team",
         italic=True, color=GREY, size=9.5)

import os
out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, "Google-IO-2026-Report.docx")
doc.save(out_path)
print("SAVED:", out_path)
